"""Document loader for tax and customs administration documents."""

from pathlib import Path
from typing import List

from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

from .config import CHUNK_OVERLAP, CHUNK_SIZE, DATA_DIR


def load_pdf(file_path: Path) -> str:
    """Load text from a PDF file."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(file_path))
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n\n".join(text_parts)
    except ImportError:
        raise ImportError("Install pypdf: pip install pypdf")
    except Exception as e:
        raise ValueError(f"Error reading PDF {file_path}: {e}")


def load_docx(file_path: Path) -> str:
    """Load text from a DOCX file."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(file_path))
        return "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
    except ImportError:
        raise ImportError("Install python-docx: pip install python-docx")
    except Exception as e:
        raise ValueError(f"Error reading DOCX {file_path}: {e}")


def load_txt(file_path: Path) -> str:
    """Load text from a plain text file."""
    encodings = ["utf-8", "utf-16", "cp1252"]
    for encoding in encodings:
        try:
            return file_path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError(f"Could not decode {file_path} with common encodings")


def load_document(file_path: Path) -> str:
    """Load a single document based on its extension."""
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return load_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return load_docx(file_path)
    elif suffix == ".txt":
        return load_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def load_documents(data_dir: Path | None = None) -> List[Document]:
    """
    Load all supported documents from the data directory.
    Supports: .pdf, .docx, .txt
    """
    data_dir = data_dir or DATA_DIR
    if not data_dir.exists():
        data_dir.mkdir(parents=True, exist_ok=True)
        return []

    documents: List[Document] = []
    supported = {".pdf", ".docx", ".doc", ".txt"}

    for file_path in data_dir.rglob("*"):
        if file_path.is_file() and file_path.suffix.lower() in supported:
            try:
                content = load_document(file_path)
                if content.strip():
                    doc = Document(
                        page_content=content,
                        metadata={"source": str(file_path.name)}
                    )
                    documents.append(doc)
            except Exception as e:
                print(f"Warning: Skipping {file_path}: {e}")

    return documents


def split_documents(documents: List[Document]) -> List[Document]:
    """Split documents into chunks for embedding."""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len,
        separators=["\n\n", "\n", " ", ""],
    )
    return text_splitter.split_documents(documents)
